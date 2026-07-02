from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).parent
ASSETS = BASE / "assets"
OUT = BASE / "OptiCrop_Project_Documentation_SkillWallet.docx"

ACCENT = "2F7D32"
ACCENT_2 = "386FA4"
MUTED = "5F6F65"
LIGHT_GREEN = "EAF4EA"
LIGHT_BLUE = "EAF2FB"
LIGHT_GOLD = "FFF3D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="1F2A24", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 26, ACCENT),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT_2),
        ("Heading 3", 11, "1F2A24"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "OptiCrop: Smart Agricultural Production Optimization Engine"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor.from_string("1F2A24")
        text = text[len(bold_lead):] if text.startswith(bold_lead) else text
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item).font.name = "Arial"


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item).font.name = "Arial"


def table(doc, headers, rows, widths=None, header_fill=LIGHT_GREEN):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9)
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=8.7)
            set_cell_margins(cells[i])
    if widths:
        for row in t.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()
    return t


def figure(doc, image_name, caption, width=6.5):
    path = ASSETS / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    for line in text.strip().splitlines():
        r = p.add_run(line.rstrip() + "\n")
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("2D333B")
    set_para_shading(p, "F6F8FA")


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def callout(doc, title, body, fill=LIGHT_GOLD):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string("1F2A24")
    p.add_run("\n" + body).font.name = "Arial"
    doc.add_paragraph()


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("OptiCrop")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Smart Agricultural Production Optimization Engine")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("1F2A24")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Complete Artificial Intelligence and Machine Learning Project Documentation")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    table(doc, ["Field", "Details"], [
        ["Project Title", "OptiCrop: Smart Agricultural Production Optimization Engine"],
        ["Team Members", "To be filled by project team"],
        ["Team Lead", "To be filled by project team lead"],
        ["Submission", "SkillWallet / SmartBridge Project Documentation"],
        ["Technologies Used", "Python 3.x, Anaconda, Jupyter/Spyder, Flask, HTML5, CSS3, Bootstrap 5, JavaScript, NumPy, Pandas, Scikit-learn, Matplotlib, Seaborn, SciPy, Pickle, Git/GitHub"],
        ["Prepared On", "June 27, 2026"],
    ], widths=[1.8, 5.5], header_fill=LIGHT_BLUE)

    callout(
        doc,
        "Project Purpose",
        "OptiCrop predicts the most suitable crop for given soil nutrient values and environmental conditions, helping stakeholders make data-driven farming decisions.",
        LIGHT_GREEN,
    )
    doc.add_page_break()


def toc(doc):
    doc.add_heading("Table of Contents", level=1)
    sections = [
        "Hardware Requirements", "Software Requirements", "Problem Statement",
        "Business Requirements", "Literature Survey", "Social and Business Impact",
        "Dataset Description", "Data Collection", "Data Preprocessing",
        "Exploratory Data Analysis", "Seasonal Crop Analysis", "Train-Test Split",
        "Machine Learning Models", "Elbow Method", "Model Evaluation",
        "Model Saving", "Flask Application", "HTML Pages", "Python Backend",
        "Application Execution", "Screenshots", "Entity Relationship Diagram",
        "Technical Architecture", "Conclusion",
    ]
    numbered(doc, sections)
    doc.add_page_break()


def requirements(doc):
    doc.add_heading("1. Hardware Requirements", level=1)
    para(doc, "The OptiCrop project is designed to run on a standard student or entry-level development machine. The model training workload is light because the dataset contains 2,200 records and seven numeric input features.")
    table(doc, ["Component", "Minimum Requirement", "Purpose"], [
        ["Processor", "Intel Core i3 or above", "Runs Jupyter Notebook, model training, and the Flask server."],
        ["RAM", "Minimum 4 GB", "Supports Python runtime, Pandas dataframes, and development tools."],
        ["Storage", "Minimum 10 GB free space", "Stores source code, dataset, notebook, virtual environment, screenshots, and model artifacts."],
        ["Internet Connection", "Required during setup", "Used for package installation, GitHub access, references, and deployment preparation."],
    ], widths=[1.4, 2.2, 3.7])

    doc.add_heading("2. Software Requirements", level=1)
    para(doc, "The software stack supports the full machine learning lifecycle: data analysis, model development, model persistence, and web deployment.")
    table(doc, ["Category", "Software / Library", "Use in OptiCrop"], [
        ["Operating System", "Windows / Linux / macOS", "Cross-platform local development and execution."],
        ["Python Runtime", "Python 3.x", "Main programming language for analysis, training, and backend routes."],
        ["IDE / Notebook", "Anaconda Navigator, Jupyter Notebook / Spyder, Visual Studio Code or PyCharm", "Interactive experimentation and source-code development."],
        ["Web Framework", "Flask Framework", "Hosts routes for Home, About, FindYourCrop, and prediction results."],
        ["Numerical Libraries", "NumPy, SciPy", "Numerical arrays, transformations, and scientific computation."],
        ["Data Libraries", "Pandas", "CSV loading, previewing, cleaning, summary statistics, and feature preparation."],
        ["ML Library", "Scikit-learn", "Train-test split, classifiers, clustering, metrics, and model evaluation."],
        ["Visualization", "Matplotlib, Seaborn", "Distribution plots, count plots, boxplots, heatmaps, and elbow graph."],
    ], widths=[1.55, 2.5, 3.25], header_fill=LIGHT_BLUE)


def business_context(doc):
    doc.add_heading("3. Problem Statement", level=1)
    para(doc, "Farmers often select crops using habit, market pressure, or incomplete local advice rather than measured soil and climate conditions. This can reduce yield, increase fertilizer waste, and create avoidable economic risk.")
    para(doc, "OptiCrop addresses this problem by learning from historical crop suitability records that contain nitrogen, phosphorous, potassium, temperature, humidity, pH, rainfall, and crop label values. The trained model receives new field observations and recommends the crop most aligned with those conditions.")
    callout(doc, "Problem Definition", "Build an AI-powered crop recommendation engine that predicts the most suitable crop based on soil nutrients and environmental parameters, then expose the recommendation through a Flask web application.", LIGHT_GREEN)

    doc.add_heading("4. Business Requirements", level=1)
    bullets(doc, [
        "Provide accurate crop recommendations from seven numeric agronomic inputs.",
        "Allow non-technical users to submit values through a simple web form.",
        "Produce reproducible analysis in a notebook and save the best model as a reusable `model.pkl` artifact.",
        "Compare multiple machine learning algorithms before selecting the final production model.",
        "Include visual analytics that explain dataset structure, feature distributions, seasonal patterns, and model performance.",
        "Keep the code modular enough for GitHub submission, classroom review, and future deployment.",
    ])
    table(doc, ["Stakeholder", "Requirement", "Expected Value"], [
        ["Farmers", "Simple recommendation interface", "Better crop choice and lower guesswork."],
        ["Agricultural Researchers", "Interpretable dataset analysis", "Evidence for crop-environment relationships."],
        ["Agribusiness Companies", "Repeatable prediction flow", "Decision support for advisory platforms."],
        ["Government / Policy Teams", "Data-driven agricultural insights", "Support for extension services and sustainable planning."],
    ], widths=[1.6, 2.8, 2.9])

    doc.add_heading("5. Literature Survey", level=1)
    para(doc, "Crop recommendation systems are a common precision-agriculture application because they translate soil and weather observations into actionable decisions. Traditional advisory systems rely on static crop calendars and expert rules, while machine learning systems can learn nonlinear relationships between nutrient levels, pH, rainfall, humidity, and crop suitability.")
    para(doc, "Classification models such as Logistic Regression, Decision Tree, Random Forest, and K-Nearest Neighbors are suitable for this problem because the target is a discrete crop label. Random Forest is often strong for tabular agricultural datasets because it handles nonlinear feature interactions, reduces single-tree overfitting, and provides stable performance without heavy feature scaling requirements.")
    para(doc, "Clustering methods such as K-Means add an unsupervised perspective by grouping similar field conditions. Although clustering does not directly predict crop labels, the resulting clusters help identify patterns in soil/environment profiles and can support exploratory seasonal or regional segmentation.")
    table(doc, ["Approach", "Description", "Relevance to OptiCrop"], [
        ["Rule-based recommendation", "Uses expert-defined thresholds for nutrients and climate.", "Useful baseline but less adaptive to complex patterns."],
        ["Supervised classification", "Learns mapping from measured features to known crop labels.", "Primary method for final prediction."],
        ["Ensemble learning", "Combines multiple decision trees or learners.", "Improves robustness and generalization."],
        ["Unsupervised clustering", "Finds natural groups in feature space.", "Supports exploratory analysis and cluster visualization."],
    ], widths=[1.9, 2.65, 2.75], header_fill=LIGHT_BLUE)

    doc.add_heading("6. Social and Business Impact", level=1)
    para(doc, "OptiCrop can improve decision quality in agriculture by giving farmers a quick second opinion based on measurable soil and environmental inputs. When used responsibly with local agronomist guidance, the tool can support better fertilizer use, water planning, and crop diversification.")
    bullets(doc, [
        "Productivity impact: recommends crops more compatible with local field conditions.",
        "Sustainability impact: reduces unnecessary fertilizer use by emphasizing nutrient-aware decisions.",
        "Water management impact: includes rainfall and humidity as model features.",
        "Economic impact: lowers risk from unsuitable crop selection and supports better planning.",
        "Educational impact: demonstrates a complete AI/ML lifecycle for SkillWallet evaluation.",
    ])


def dataset_and_preprocessing(doc):
    doc.add_heading("7. Dataset Description", level=1)
    para(doc, "The project uses `Crop_recommendation.csv`, a public tabular crop recommendation dataset commonly distributed through Kaggle as the Crop Recommendation Dataset. It contains 2,200 records, eight columns, seven numeric features, and one target label.")
    table(doc, ["Column", "Type", "Description"], [
        ["N", "Integer", "Nitrogen ratio/content in soil."],
        ["P", "Integer", "Phosphorous ratio/content in soil."],
        ["K", "Integer", "Potassium ratio/content in soil."],
        ["temperature", "Float", "Ambient temperature observed for crop conditions."],
        ["humidity", "Float", "Relative humidity percentage."],
        ["ph", "Float", "Soil pH value indicating acidity or alkalinity."],
        ["rainfall", "Float", "Rainfall measurement associated with crop conditions."],
        ["label", "Object / Category", "Crop name to be predicted, such as rice, maize, banana, cotton, jute, coffee, and others."],
    ], widths=[1.3, 1.5, 4.5])
    figure(doc, "dataset_preview.png", "Figure 1: Dataset preview showing the first records from Crop_recommendation.csv.", 6.6)
    para(doc, "Dataset source: Kaggle Crop Recommendation Dataset, commonly available at `https://www.kaggle.com/datasets/atharvaingle/crop-recommendation-dataset`.")

    doc.add_heading("8. Data Collection", level=1)
    para(doc, "Data collection begins by downloading the CSV file and placing it in the project structure under `dataset/Crop_recommendation.csv`. The notebook loads the file with `pd.read_csv()` and validates the columns before analysis.")
    code_block(doc, """
import pandas as pd
data = pd.read_csv("dataset/Crop_recommendation.csv")
data.head()
""")
    para(doc, "The collected fields represent both soil fertility indicators and climate/environment variables. This makes the dataset suitable for a crop recommendation model because the prediction is based on the same factors farmers and agronomists evaluate in field planning.")

    doc.add_heading("9. Data Preprocessing", level=1)
    para(doc, "Data preprocessing converts the raw CSV into reliable training input. The process checks shape, column data types, missing values, duplicate values, numerical summaries, and potential outliers.")
    table(doc, ["Preprocessing Step", "Command / Method", "Purpose"], [
        ["Dataset loading", "`pd.read_csv()`", "Loads the CSV into a Pandas dataframe."],
        ["Dataset preview", "`head()`", "Confirms columns and example values."],
        ["Shape", "`shape`", "Confirms expected 2,200 records and eight columns."],
        ["Data types", "`info()`", "Checks numeric feature types and non-null counts."],
        ["Summary statistics", "`describe()`", "Measures count, mean, standard deviation, quartiles, and ranges."],
        ["Missing values", "`isnull().sum()`", "Ensures no absent values remain before training."],
        ["Outlier detection", "IQR method", "Identifies values beyond lower/upper bounds."],
        ["Transformation", "Log transformation where required", "Reduces skew for highly skewed positive features such as rainfall."],
    ], widths=[1.8, 2.2, 3.3], header_fill=LIGHT_BLUE)
    figure(doc, "dataset_info.png", "Figure 2: Dataset info() evidence showing non-null counts and data types.", 6.6)
    figure(doc, "describe_summary.png", "Figure 3: describe() statistical summary for numeric features.", 6.6)
    figure(doc, "missing_values.png", "Figure 4: Missing values check confirming clean input columns.", 6.6)

    doc.add_heading("Outlier Detection and IQR", level=2)
    para(doc, "The Interquartile Range method is used to detect extreme values. Q1 is the 25th percentile and Q3 is the 75th percentile. The IQR is Q3 minus Q1.")
    code_block(doc, """
Q1 = data[numeric_columns].quantile(0.25)
Q3 = data[numeric_columns].quantile(0.75)
IQR = Q3 - Q1
lower_bound = Q1 - 1.5 * IQR
upper_bound = Q3 + 1.5 * IQR
""")
    para(doc, "Values outside the lower and upper bounds are reviewed. In agricultural data, outliers are not always errors; high rainfall or unusual nutrient values can represent real crop-specific conditions. Therefore, OptiCrop detects outliers, visualizes them, and applies log transformation only when skewness makes modeling or interpretation harder.")
    figure(doc, "boxplot_iqr.png", "Figure 5: Boxplot evidence for outlier detection using IQR logic.", 6.4)


def eda_and_split(doc):
    doc.add_heading("10. Exploratory Data Analysis", level=1)
    para(doc, "Exploratory Data Analysis helps understand the relationship between features and crop labels before model training. OptiCrop uses univariate, bivariate, and multivariate analysis.")
    table(doc, ["EDA Type", "Techniques", "Interpretation"], [
        ["Univariate Analysis", "Histograms, KDE/distribution plots, count plots", "Shows feature spread, skewness, and label balance."],
        ["Bivariate Analysis", "Scatter plots, grouped summaries, boxplots by crop", "Shows how one feature relates to the target or another feature."],
        ["Multivariate Analysis", "Correlation heatmap and pairwise feature views", "Reveals relationships among multiple numeric variables."],
    ], widths=[1.7, 2.75, 2.85])
    figure(doc, "distribution_plots.png", "Figure 6: Distribution plots for nutrient and environmental features.", 6.7)
    figure(doc, "count_plot.png", "Figure 7: Count plot showing balanced crop-label records.", 6.7)
    figure(doc, "correlation_heatmap.png", "Figure 8: Multivariate correlation heatmap for numeric crop conditions.", 5.7)
    figure(doc, "source_page-02.png", "Figure 9: Source screenshot page showing library imports and distribution-plot evidence.", 5.2)

    doc.add_heading("11. Seasonal Crop Analysis", level=1)
    para(doc, "Seasonal crop analysis groups labels according to climatic suitability. This is useful because a crop recommendation should be interpreted with seasonal availability, rainfall, and temperature patterns rather than only raw model output.")
    table(doc, ["Season", "Example Crops", "Typical Interpretation"], [
        ["Summer", "maize, cotton, mango, muskmelon, watermelon", "Higher temperature tolerance and irrigation/rainfall planning are important."],
        ["Winter", "chickpea, lentil, kidneybeans, orange, apple", "Cooler conditions and lower humidity may be more suitable."],
        ["Rainy", "rice, jute, coconut, papaya, pigeonpeas", "Rainfall and humidity are stronger decision factors."],
    ], widths=[1.3, 3.0, 3.0], header_fill=LIGHT_GREEN)
    figure(doc, "seasonal_crop_analysis.png", "Figure 10: Seasonal crop analysis visualization.", 5.8)

    doc.add_heading("12. Train-Test Split", level=1)
    para(doc, "The dataset is separated into features `X` and target `y`. The seven input features are nitrogen, phosphorous, potassium, temperature, humidity, pH, and rainfall. The target is the crop `label`.")
    code_block(doc, """
X = data.drop("label", axis=1)
y = data["label"]

from sklearn.model_selection import train_test_split
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=42, stratify=y
)
""")
    para(doc, "The recommended split is 80% training and 20% testing. Stratification keeps the crop-label distribution similar in both subsets, which is important because each crop class has a balanced number of records.")
    table(doc, ["Subset", "Percentage", "Approximate Records", "Purpose"], [
        ["Training", "80%", "1,760", "Fit model parameters and learn crop patterns."],
        ["Testing", "20%", "440", "Evaluate performance on unseen examples."],
    ], widths=[1.5, 1.3, 1.8, 2.7], header_fill=LIGHT_BLUE)


def models_and_eval(doc):
    doc.add_heading("13. Machine Learning Models", level=1)
    para(doc, "OptiCrop trains and compares multiple algorithms to avoid selecting a model blindly. Each model has a different bias, interpretability level, and sensitivity to feature scale.")
    table(doc, ["Model", "How It Works", "Strength in OptiCrop", "Caution"], [
        ["KNN", "Predicts using the nearest training records in feature space.", "Simple and intuitive for similar field-condition matching.", "Sensitive to feature scaling and distance metrics."],
        ["Logistic Regression", "Learns linear decision boundaries for multiclass labels.", "Fast baseline and interpretable coefficients.", "May underfit nonlinear crop-environment relationships."],
        ["Decision Tree", "Splits features into rule-like branches.", "Highly interpretable and handles nonlinear thresholds.", "Can overfit if tree depth is not controlled."],
        ["Random Forest", "Combines many decision trees.", "Strong accuracy and stable generalization for tabular data.", "Less transparent than a single tree."],
        ["K-Means Clustering", "Groups records by feature similarity without labels.", "Useful for pattern discovery and cluster analysis.", "Not a direct supervised predictor."],
    ], widths=[1.25, 2.2, 2.1, 1.75])
    figure(doc, "model_comparison.png", "Figure 11: Model comparison showing Random Forest as the strongest candidate.", 5.9)

    doc.add_heading("14. Elbow Method", level=1)
    para(doc, "For K-Means clustering, WCSS means Within-Cluster Sum of Squares. It measures how compact the clusters are by summing squared distances between each point and its assigned cluster centroid.")
    para(doc, "As `k` increases, WCSS decreases because more centroids can represent the data. The optimal number of clusters is usually near the elbow point, where WCSS reduction begins to slow. In OptiCrop, the elbow graph supports choosing a practical cluster count rather than overfitting many tiny clusters.")
    figure(doc, "elbow_graph.png", "Figure 12: Elbow graph with WCSS and the optimal-cluster bend.", 6.1)

    doc.add_heading("15. Model Evaluation", level=1)
    para(doc, "Model evaluation uses accuracy, precision, recall, F1 score, confusion matrix, and classification report. These metrics verify not only overall correctness but also class-level behavior.")
    table(doc, ["Metric", "Meaning", "Why It Matters"], [
        ["Accuracy", "Correct predictions divided by total predictions.", "Gives overall performance."],
        ["Precision", "Correct positive predictions divided by all predicted positives.", "Checks how reliable predicted crop labels are."],
        ["Recall", "Correct positive predictions divided by actual positives.", "Checks whether each crop class is being captured."],
        ["F1 Score", "Harmonic mean of precision and recall.", "Balances precision and recall for class-level evaluation."],
        ["Confusion Matrix", "Table of actual vs predicted labels.", "Shows exactly which crops are confused."],
        ["Classification Report", "Precision, recall, F1, and support for every class.", "Provides detailed model validation for submission."],
    ], widths=[1.35, 2.55, 3.4], header_fill=LIGHT_BLUE)
    figure(doc, "classification_report.png", "Figure 13: Classification report evidence for final model performance.", 6.6)
    para(doc, "The final model should be selected using test-set performance and practical reliability. For this dataset, Random Forest is commonly selected because it provides high accuracy on multiclass tabular crop recommendation data and is robust to nonlinear relationships.")


def deployment(doc):
    doc.add_heading("16. Model Saving", level=1)
    para(doc, "After the best model is selected, it is serialized using Python Pickle. The saved artifact is named `model.pkl` and is loaded by the Flask application at runtime.")
    code_block(doc, """
import pickle

with open("model.pkl", "wb") as file:
    pickle.dump(best_model, file)

with open("model.pkl", "rb") as file:
    model = pickle.load(file)
""")
    para(doc, "Saving the model separates training from deployment. The web application does not need to retrain the model each time a user opens the page; it only loads the saved model and performs prediction.")

    doc.add_heading("17. Flask Application", level=1)
    para(doc, "The Flask layer turns the trained model into an interactive web application. It receives form inputs from the user, validates numeric values, converts them into the expected model format, calls the prediction function, and renders the result page.")
    table(doc, ["Route", "Page / Function", "Responsibility"], [
        ["/", "Home", "Displays navigation, project overview, hero section, and footer."],
        ["/about", "About", "Explains objectives, technologies, ML workflow, benefits, and scope."],
        ["/findyourcrop", "FindYourCrop", "Displays form fields for N, P, K, temperature, humidity, pH, and rainfall."],
        ["/predict", "Prediction Result", "Handles POST request, loads values, calls model prediction, and returns crop result."],
    ], widths=[1.2, 2.0, 4.1])
    figure(doc, "source_page-10.png", "Figure 14: Source screenshot page showing HTML navigation and prediction-form code evidence.", 5.3)

    doc.add_heading("18. HTML Pages", level=1)
    para(doc, "The HTML layer contains separate templates for Home, About, FindYourCrop, and Result. Bootstrap 5 is recommended for responsiveness and consistent styling.")
    table(doc, ["Template", "Main Content", "Expected Controls / Visuals"], [
        ["home.html", "Navigation bar, hero section, project overview, agriculture background image, footer.", "Home/About/FindYourCrop links and call-to-action button."],
        ["about.html", "Project description, objectives, technologies, workflow, benefits, scope.", "Readable sections and cards for project value."],
        ["findyourcrop.html", "Input form for seven agronomic values.", "Numeric fields and Predict button."],
        ["result.html", "Predicted crop and confirmation message.", "Result card and link back to the form."],
    ], widths=[1.4, 3.4, 2.5], header_fill=LIGHT_BLUE)
    figure(doc, "home_page.png", "Figure 15: Home Page screenshot.", 6.4)
    figure(doc, "about_page.png", "Figure 16: About Page screenshot.", 6.4)
    figure(doc, "findyourcrop_page.png", "Figure 17: FindYourCrop Page screenshot.", 6.4)
    figure(doc, "prediction_result.png", "Figure 18: Prediction Result screenshot.", 6.4)

    doc.add_heading("19. Python Backend", level=1)
    para(doc, "The backend begins by initializing Flask, loading the saved Pickle model, defining route functions, extracting POST values, converting inputs to floats, calling `model.predict()`, and returning the rendered output template.")
    code_block(doc, """
from flask import Flask, render_template, request
import pickle
import numpy as np

app = Flask(__name__)
model = pickle.load(open("model.pkl", "rb"))

@app.route("/")
def home():
    return render_template("home.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/findyourcrop")
def findyourcrop():
    return render_template("findyourcrop.html")

@app.route("/predict", methods=["POST"])
def predict():
    values = [float(request.form[field]) for field in
              ["N", "P", "K", "temperature", "humidity", "ph", "rainfall"]]
    prediction = model.predict(np.array(values).reshape(1, -1))[0]
    return render_template("result.html", prediction=prediction)

if __name__ == "__main__":
    app.run(debug=True)
""")

    doc.add_heading("20. Application Execution", level=1)
    numbered(doc, [
        "Open the OptiCrop project folder in VS Code, PyCharm, or terminal.",
        "Create or activate a Python environment with the required libraries installed.",
        "Ensure `model.pkl` exists in the expected path after training.",
        "Run `python app.py`.",
        "Open `http://127.0.0.1:5000` in the browser.",
        "Navigate to FindYourCrop, enter values, click Predict, and verify the result page.",
    ])
    code_block(doc, """
pip install -r requirements.txt
python model/train_model.py
python app.py
""")


def screenshots_section(doc):
    doc.add_heading("21. Screenshots", level=1)
    para(doc, "This section consolidates the required screenshot evidence. Generated charts and UI captures are placed under their relevant sections above; the supplied PDF evidence pages are also included below as source screenshots.")
    table(doc, ["Required Screenshot", "Included As"], [
        ["Dataset", "Figure 1 and source page 1"],
        ["Graphs / Distribution plots", "Figure 6 and source page 2"],
        ["Boxplot", "Figure 5"],
        ["Elbow Graph", "Figure 12"],
        ["Classification Report", "Figure 13"],
        ["Prediction Result", "Figure 18"],
        ["Home Page", "Figure 15"],
        ["About Page", "Figure 16"],
        ["FindYourCrop Page", "Figure 17"],
    ], widths=[2.6, 4.7], header_fill=LIGHT_GREEN)
    for i in range(1, 11):
        figure(doc, f"source_page-{i:02}.png", f"Source Screenshot Page {i}: Evidence extracted from the supplied OptiCrop PDF.", 4.95)


def architecture_and_close(doc):
    doc.add_heading("22. Entity Relationship Diagram", level=1)
    para(doc, "OptiCrop primarily uses a CSV dataset and Pickle model artifact rather than a mandatory relational database. The ERD below models the logical entities that would exist if prediction history were stored for production use.")
    figure(doc, "entity_relationship_diagram.png", "Figure 19: Logical ERD for storing inputs, predictions, crop labels, and model artifacts.", 6.7)
    para(doc, "The `CropInput` entity stores user-submitted feature values. The `Prediction` entity records the predicted crop returned by the model. The `CropLabel` entity stores metadata about possible crop classes, and `ModelArtifact` tracks the deployed model version.")

    doc.add_heading("23. Technical Architecture", level=1)
    para(doc, "The technical architecture separates data science, model persistence, and web serving. The notebook is used for experimentation and evaluation. The selected model is saved as `model.pkl`. Flask loads that artifact and connects it with HTML templates for user interaction.")
    figure(doc, "technical_architecture.png", "Figure 20: Technical architecture of the OptiCrop application.", 6.7)
    table(doc, ["Layer", "Components", "Responsibility"], [
        ["Presentation Layer", "HTML5, CSS3, Bootstrap 5, JavaScript", "Collect inputs and display crop recommendation results."],
        ["Application Layer", "Flask routes in `app.py`", "Handle navigation, POST requests, validation, and template rendering."],
        ["ML Layer", "Scikit-learn model loaded from `model.pkl`", "Predict crop label from feature array."],
        ["Data Layer", "`Crop_recommendation.csv` and notebook outputs", "Provide training data, analysis, charts, and evaluation evidence."],
    ], widths=[1.6, 2.5, 3.2], header_fill=LIGHT_BLUE)

    doc.add_heading("24. Conclusion", level=1)
    para(doc, "OptiCrop demonstrates a complete end-to-end AI and machine learning workflow for agricultural production optimization. The project starts with a clearly defined farming decision problem, uses a structured crop recommendation dataset, performs preprocessing and exploratory analysis, trains multiple models, evaluates them with classification metrics, saves the best model, and deploys it through a Flask web interface.")
    para(doc, "The project is suitable for SkillWallet submission because it includes the required hardware and software specifications, business context, literature survey, dataset description, preprocessing workflow, EDA, seasonal analysis, supervised and unsupervised learning, model evaluation, model saving, web application structure, screenshots, ERD, architecture, and conclusion.")
    callout(doc, "Final Outcome", "The completed OptiCrop documentation presents the project as a professional, reproducible, and submission-ready AI/ML application for smart crop recommendation.", LIGHT_GREEN)


def main():
    doc = Document()
    style_document(doc)
    cover(doc)
    toc(doc)
    requirements(doc)
    business_context(doc)
    dataset_and_preprocessing(doc)
    eda_and_split(doc)
    models_and_eval(doc)
    deployment(doc)
    screenshots_section(doc)
    architecture_and_close(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
